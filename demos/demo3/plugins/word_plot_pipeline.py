"""Plugin: folder-enter → write to Word; file-read → collect data; folder-exit → plot and paste into Word.

Dependencies: matplotlib, python-docx. The processors handle missing dependencies gracefully.
Usage: reference these processor names in config rules via pre_processors/processors/post_processors.
"""
from pathlib import Path
from typing import List, Dict, Any
from decorators.processor import processor
from utils.pipeline import (
    ensure_dir,
    get_bucket,
    append_numbers,
    get_or_create_doc,
    set_output,
)
from utils.adapters.docx_helpers import get_or_create_doc as adapter_get_doc
import os
import time

# 强制使用非交互后端，需在任何 matplotlib 导入前设置
os.environ.setdefault("MPLBACKEND", "Agg")


def _ensure_packages():
    """Ensure required packages; do NOT import pyplot to avoid GUI backend activation.

    Uses Figure + Agg canvas explicitly later. This prevents thread-related GUI crashes.
    """
    try:
        import matplotlib  # noqa: F401
        from docx import Document  # noqa: F401
        from matplotlib.figure import Figure  # noqa: F401
        from matplotlib.backends.backend_agg import FigureCanvasAgg  # noqa: F401
        return True, None
    except Exception as e:
        return False, e


@processor(name="enter_dir_write_word", priority=80, source=__file__, metadata={
    "name": "Enter Dir → Write Word",
    "author": "pipeline",
    "version": "0.1",
    "description": "When entering a directory, write a header to a Word document",
})
def enter_dir_write_word(path: Path, context, **cfg):
    if not path.is_dir():
        return {"skipped": True}
    ok, err = _ensure_packages()
    if not ok:
        return {"error": f"missing packages: {err}"}
    from docx import Document

    out_doc = Path(cfg.get("doc_path", "./output.docx"))
    doc = Document(out_doc) if out_doc.exists() else Document()
    doc.add_heading(f"进入目录: {path.name}", level=2)
    doc.add_paragraph(f"路径: {path}")
    doc.save(out_doc)
    # track doc path in shared
    context.set_shared(["word_doc"], str(out_doc))
    set_output(context, "enter_dir", "doc_path", str(out_doc))
    return {"doc": str(out_doc), "action": "write_enter"}


@processor(name="read_data_files", priority=70, source=__file__, metadata={
    "name": "Read Data Files",
    "author": "pipeline",
    "version": "0.1",
    "description": "Read simple numeric data from text/CSV files under a directory",
})
def read_data_files(path: Path, context, **cfg):
    # If called on a directory, read all files matching pattern; if on a file, just read that file.
    import csv, json, re
    pattern = cfg.get("pattern", "*")
    key = cfg.get("key", "values")
    collected: List[float] = get_bucket(context, ["folder_data", str(path)], [])

    def read_one(file_path: Path):
        vals = []
        try:
            text = file_path.read_text(encoding="utf-8", errors="ignore")
            # CSV
            if file_path.suffix == ".csv" or text.strip().startswith("category,") or "," in text.splitlines()[0]:
                reader = csv.reader(text.splitlines())
                header = next(reader, None)
                for row in reader:
                    for v in row:
                        try:
                            vals.append(float(v))
                        except Exception:
                            pass
                return vals
            # JSON
            if text.strip().startswith("[") or text.strip().startswith("{"):
                try:
                    data = json.loads(text)
                    # 支持列表或字典
                    if isinstance(data, list):
                        for item in data:
                            if isinstance(item, dict):
                                for k, v in item.items():
                                    if isinstance(v, (int, float)):
                                        vals.append(v)
                                    elif isinstance(v, dict):
                                        for vv in v.values():
                                            if isinstance(vv, (int, float)):
                                                vals.append(vv)
                    elif isinstance(data, dict):
                        for v in data.values():
                            if isinstance(v, (int, float)):
                                vals.append(v)
                            elif isinstance(v, list):
                                for vv in v:
                                    if isinstance(vv, (int, float)):
                                        vals.append(vv)
                except Exception:
                    pass
                return vals
            # 混合格式
            if "CSV:" in text or "JSON:" in text:
                for line in text.splitlines():
                    line = line.strip()
                    if line.startswith("CSV:"):
                        continue
                    if re.match(r"^[A-Za-z]+,\d+", line):
                        parts = line.split(",")
                        for v in parts[1:]:
                            try:
                                vals.append(float(v))
                            except Exception:
                                pass
                    if line.startswith("JSON:"):
                        try:
                            j = json.loads(line[5:].strip())
                            if isinstance(j, dict):
                                for v in j.values():
                                    if isinstance(v, (int, float)):
                                        vals.append(v)
                        except Exception:
                            pass
                    if re.match(r"^\d+$", line):
                        vals.append(float(line))
                return vals
            # TXT: 多列数字
            for line in text.splitlines():
                line = line.strip()
                if not line or line.startswith("#"):
                    continue
                # 逗号或空格分隔
                for part in re.split(r",|\s+", line):
                    try:
                        vals.append(float(part))
                    except Exception:
                        pass
        except Exception:
            return []
        return vals

    files = []
    if path.is_dir():
        files = list(path.glob(pattern))
    elif path.is_file():
        files = [path]

    added = 0
    for f in files:
        vals = read_one(f)
        if not vals:
            continue
        # 将文件数据归并到所属目录键，确保目录退出时可绘图
        if f.is_file():
            dir_key = ["folder_data", str(f.parent)]
            append_numbers(context, dir_key, vals)
        collected.extend(vals)
        added += len(vals)

    return {"path": str(path), key: len(collected), "added": added, "files": [str(f) for f in files]}


@processor(name="plot_on_exit_paste_word", priority=60, source=__file__, metadata={
    "name": "Plot on Exit → Paste Word",
    "author": "pipeline",
    "version": "0.1",
    "description": "On directory exit: plot collected values and paste into Word",
})
def plot_on_exit_paste_word(path: Path, context, **cfg):
    if not path.is_dir():
        return {"skipped": True}
    ok, err = _ensure_packages()
    if not ok:
        return {"error": f"missing packages: {err}"}
    # 调试日志文件（写入磁盘便于分析闪退前阶段）
    debug_dir = ensure_dir(Path(cfg.get("debug_dir", "./debug_logs")))
    log_file = debug_dir / f"debug_{path.name}.log"
    def _log_step(msg: str):
        try:
            with open(log_file, 'a', encoding='utf-8') as lf:
                lf.write(f"[{time.strftime('%H:%M:%S')}] {msg}\n")
        except Exception:
            pass

    _log_step("BEGIN plot_on_exit_paste_word")
    _log_step(f"Values count: {len(context.get_data(['folder_data', str(path)], []))}")

    # 先准备数据
    from docx import Document
    from docx.shared import Inches

    values = context.get_data(["folder_data", str(path)], [])
    if not values:
        _log_step("No data available; writing placeholder")
        # 写入占位“离开目录”说明（用户希望看到离开信息）
        doc_path_placeholder = Path(context.get_shared(["word_doc"], cfg.get("doc_path", "./output.docx"))).resolve()
        try:
            from docx import Document
            doc_ph = Document(doc_path_placeholder) if doc_path_placeholder.exists() else Document()
            doc_ph.add_heading(f"离开目录: {path.name}", level=3)
            doc_ph.add_paragraph("无数据可绘图。")
            doc_ph.save(doc_path_placeholder)
        except Exception as e:
            _log_step(f"Placeholder write failed: {e}")
        return {"path": str(path), "info": "no data", "doc": str(doc_path_placeholder)}

    # 选择绘图方式：优先 matplotlib，失败则 Pillow 后备
    use_pillow = cfg.get("force_pillow", False)
    img_dir = Path(cfg.get("img_dir", "./images"))
    img_dir.mkdir(parents=True, exist_ok=True)
    img_path = (img_dir / f"plot_{path.name}.png").resolve()

    def _render_with_matplotlib():
        from matplotlib.figure import Figure
        from matplotlib.backends.backend_agg import FigureCanvasAgg as FigureCanvas
        fig = Figure(figsize=(cfg.get("fig_width", 4), cfg.get("fig_height", 3)), dpi=cfg.get("dpi", 100))
        ax = fig.add_subplot(111)
        ax.plot(values, marker="o")
        ax.set_title(f"数据曲线: {path.name}")
        ax.grid(True)
        fig.tight_layout()
        canvas = FigureCanvas(fig)
        canvas.draw()
        fig.savefig(img_path)
        return True

    def _render_with_pillow():
        try:
            from PIL import Image, ImageDraw, ImageFont
        except Exception as e:
            _log_step(f"Pillow import failed: {e}")
            return False
        W, H = int(cfg.get("px_width", 600)), int(cfg.get("px_height", 400))
        im = Image.new("RGB", (W, H), (255, 255, 255))
        draw = ImageDraw.Draw(im)
        margin = 40
        plot_w = W - 2 * margin
        plot_h = H - 2 * margin
        if not values:
            draw.text((margin, margin), "NO DATA", fill=(0, 0, 0))
        else:
            vmin, vmax = min(values), max(values)
            rng = vmax - vmin if vmax != vmin else 1.0
            pts = []
            for i, v in enumerate(values):
                x = margin + i / max(1, len(values) - 1) * plot_w
                y = margin + (1 - (v - vmin) / rng) * plot_h
                pts.append((x, y))
            # 轴
            draw.rectangle([margin, margin, margin + plot_w, margin + plot_h], outline=(0, 0, 0))
            if len(pts) > 1:
                draw.line(pts, fill=(30, 120, 200), width=2)
            for x, y in pts:
                draw.ellipse([x - 3, y - 3, x + 3, y + 3], fill=(200, 50, 50))
            title = f"数据曲线: {path.name}"
            try:
                font = ImageFont.load_default()
                draw.text((margin, 5), title, fill=(0, 0, 0), font=font)
            except Exception:
                draw.text((margin, 5), title, fill=(0, 0, 0))
        im.save(img_path)
        return True

    render_ok = False
    if not use_pillow:
        try:
            _log_step("Attempt matplotlib render")
            render_ok = _render_with_matplotlib()
        except Exception as e:
            _log_step(f"Matplotlib render failed: {e}; fallback Pillow")
            use_pillow = True
    if use_pillow and not render_ok:
        render_ok = _render_with_pillow()

    if not render_ok or not img_path.exists():
        _log_step("Image render failed")
        return {"path": str(path), "error": "image_not_rendered", "img": str(img_path)}

    _log_step("Image rendered successfully")

    # paste into Word
    if not img_path.exists():
        _log_step("Image file missing after render")
        return {"path": str(path), "error": f"image not created: {img_path}"}

    doc_path_cfg = context.get_shared(["word_doc"], cfg.get("doc_path", "./output.docx"))
    # record path in context, then use adapter to open Document
    doc_path = get_or_create_doc(context, doc_path_cfg)
    from utils.adapters.docx_helpers import get_or_create_doc as _open_doc
    doc, doc_path = _open_doc(str(doc_path))
    doc.add_heading(f"离开目录: {path.name}", level=3)
    doc.add_paragraph(f"文件数: {len(values)}；示例曲线如下：")
    try:
        width_in = cfg.get("img_width_inch", 4.0)
        doc.add_picture(str(img_path), width=Inches(width_in))
        _log_step("Picture added to document")
    except Exception as e:
        _log_step(f"Failed to add picture: {e}")
        return {"doc": str(doc_path), "image": str(img_path), "action": "plot_paste_exit", "error": f"failed to add picture: {e}"}
    doc.save(doc_path)
    _log_step("Document saved successfully")
    set_output(context, "plot_exit", "image_path", str(img_path))
    return {"doc": str(doc_path), "image": str(img_path), "action": "plot_paste_exit", "debug_log": str(log_file)}
