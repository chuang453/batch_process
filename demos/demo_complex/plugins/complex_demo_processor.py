"""Complex demo plugin: handles folder labels, 3-type data reading, plotting, and Word document generation.

This plugin demonstrates:
1. Writing folder labels when entering directories
2. Reading files with 3 types of data (Type1: main data occupies 2 cells, Type2 & Type3: 1 cell each)
3. Writing data to Word tables
4. Creating individual plots for each file
5. Writing summary and comprehensive plot when exiting directories

Dependencies: matplotlib, python-docx, Pillow
"""
from pathlib import Path
from typing import List, Dict, Any, Tuple
from decorators.processor import processor
from utils.pipeline import ensure_dir, get_bucket, set_output
import os
import time
import re

# Force non-interactive backend before any matplotlib import
os.environ.setdefault("MPLBACKEND", "Agg")


def _ensure_packages():
    """Ensure required packages are available."""
    try:
        import matplotlib  # noqa: F401
        from docx import Document  # noqa: F401
        from matplotlib.figure import Figure  # noqa: F401
        from matplotlib.backends.backend_agg import FigureCanvasAgg  # noqa: F401
        from PIL import Image  # noqa: F401
        return True, None
    except Exception as e:
        return False, e


# Folder labels mapping
FOLDER_LABELS = {
    "folder_A": "实验组A - 温度控制实验",
    "folder_B": "实验组B - 电力测试实验",
    "folder_C": "实验组C - 化学分析实验",
    "subfolder_A1": "子组A1 - 高温条件",
    "subfolder_A2": "子组A2 - 低温条件",
    "subfolder_B1": "子组B1 - 标准电压",
    "subfolder_B2": "子组B2 - 变压测试",
    "subfolder_C1": "子组C1 - 酸碱度测试",
}


def parse_three_type_data(file_path: Path) -> Tuple[List[float], List[float], List[float]]:
    """Parse file with 3 types of data.
    
    Returns:
        Tuple of (type1_data, type2_data, type3_data)
    """
    type1_data = []
    type2_data = []
    type3_data = []
    
    try:
        text = file_path.read_text(encoding="utf-8", errors="ignore")
        for line in text.splitlines():
            line = line.strip()
            if not line or line.startswith("#"):
                continue
            
            if line.startswith("TYPE1:"):
                # Extract numbers after TYPE1:
                nums_str = line[6:].strip()
                for num in re.split(r',\s*', nums_str):
                    try:
                        type1_data.append(float(num))
                    except:
                        pass
            elif line.startswith("TYPE2:"):
                nums_str = line[6:].strip()
                for num in re.split(r',\s*', nums_str):
                    try:
                        type2_data.append(float(num))
                    except:
                        pass
            elif line.startswith("TYPE3:"):
                nums_str = line[6:].strip()
                for num in re.split(r',\s*', nums_str):
                    try:
                        type3_data.append(float(num))
                    except:
                        pass
    except Exception as e:
        print(f"Error parsing file {file_path}: {e}")
    
    return type1_data, type2_data, type3_data


@processor(name="enter_folder_label", priority=90, source=__file__, metadata={
    "name": "Enter Folder - Write Label",
    "author": "complex_demo",
    "version": "1.0",
    "description": "When entering a folder, write its label to Word document",
})
def enter_folder_label(path: Path, context, **cfg):
    """Write folder label when entering a directory."""
    if not path.is_dir():
        return {"skipped": True}
    
    ok, err = _ensure_packages()
    if not ok:
        return {"error": f"missing packages: {err}"}
    
    from docx import Document
    from docx.shared import Pt, RGBColor
    
    # Get or create document
    out_doc = Path(cfg.get("doc_path", "./demo_complex_output.docx"))
    doc = Document(out_doc) if out_doc.exists() else Document()
    
    # Get folder label
    folder_name = path.name
    label = FOLDER_LABELS.get(folder_name, folder_name)
    
    # Add heading with label
    heading = doc.add_heading(f"📁 {label}", level=2)
    heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
    
    # Add path info
    doc.add_paragraph(f"路径: {path.relative_to(context.root_path)}")
    doc.add_paragraph("")  # blank line
    
    doc.save(out_doc)
    
    # Store doc path in context
    context.set_shared(["complex_demo", "doc_path"], str(out_doc))
    
    return {
        "doc": str(out_doc),
        "folder": folder_name,
        "label": label,
        "action": "enter_folder_label"
    }


@processor(name="read_three_type_data", priority=70, source=__file__, metadata={
    "name": "Read 3-Type Data",
    "author": "complex_demo",
    "version": "1.0",
    "description": "Read files with 3 types of data, write to Word table, and create plot",
})
def read_three_type_data(path: Path, context, **cfg):
    """Read file with 3 types of data, write to Word, and create individual plot."""
    if not path.is_file():
        return {"skipped": True}
    
    ok, err = _ensure_packages()
    if not ok:
        return {"error": f"missing packages: {err}"}
    
    from docx import Document
    from docx.shared import Inches, Pt
    from matplotlib.figure import Figure
    from matplotlib.backends.backend_agg import FigureCanvasAgg as FigureCanvas
    
    # Parse data
    type1, type2, type3 = parse_three_type_data(path)
    
    if not (type1 or type2 or type3):
        return {"skipped": True, "reason": "no_data"}
    
    # Get document
    doc_path = context.get_shared(["complex_demo", "doc_path"], cfg.get("doc_path", "./demo_complex_output.docx"))
    doc = Document(doc_path) if Path(doc_path).exists() else Document()
    
    # Add file name as subheading
    doc.add_heading(f"文件: {path.name}", level=3)
    
    # Create table with data
    # Format: Type1 occupies 2 columns, Type2 and Type3 occupy 1 column each (total 4 columns)
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Type1 主要数据"
    header_cells[0].merge(header_cells[1])  # Merge 2 cells for Type1
    header_cells[2].text = "Type2 辅助数据1"
    header_cells[3].text = "Type3 辅助数据2"
    
    # Data row
    data_cells = table.rows[1].cells
    type1_str = ", ".join(f"{v:.2f}" for v in type1[:10])  # Show first 10 values
    if len(type1) > 10:
        type1_str += f" ... (共{len(type1)}个)"
    data_cells[0].text = type1_str
    data_cells[0].merge(data_cells[1])  # Merge 2 cells for Type1
    
    type2_str = ", ".join(f"{v:.2f}" for v in type2[:5])
    if len(type2) > 5:
        type2_str += f" ... (共{len(type2)}个)"
    data_cells[2].text = type2_str
    
    type3_str = ", ".join(f"{v:.2f}" for v in type3[:5])
    if len(type3) > 5:
        type3_str += f" ... (共{len(type3)}个)"
    data_cells[3].text = type3_str
    
    # Create individual plot
    img_dir = Path(cfg.get("img_dir", "./demo_complex_images"))
    img_dir.mkdir(parents=True, exist_ok=True)
    img_path = img_dir / f"plot_{path.stem}.png"
    
    try:
        fig = Figure(figsize=(6, 4), dpi=100)
        ax = fig.add_subplot(111)
        
        if type1:
            ax.plot(type1, marker='o', label='Type1 主要数据', linewidth=2, markersize=4)
        if type2:
            ax.plot(type2, marker='s', label='Type2 辅助数据1', linewidth=1.5, markersize=4)
        if type3:
            ax.plot(type3, marker='^', label='Type3 辅助数据2', linewidth=1.5, markersize=4)
        
        ax.set_title(f"数据曲线 - {path.stem}")
        ax.set_xlabel("索引")
        ax.set_ylabel("数值")
        ax.legend()
        ax.grid(True, alpha=0.3)
        fig.tight_layout()
        
        canvas = FigureCanvas(fig)
        canvas.draw()
        fig.savefig(img_path)
        
        # Add plot to Word
        doc.add_paragraph("数据可视化：")
        doc.add_picture(str(img_path), width=Inches(5.5))
        doc.add_paragraph("")  # blank line
        
    except Exception as e:
        doc.add_paragraph(f"绘图失败: {e}")
    
    doc.save(doc_path)
    
    # Store data for folder summary
    folder_key = str(path.parent)
    folder_bucket = context.setdefault_data(["complex_demo", "folder_data", folder_key], {
        "type1": [], "type2": [], "type3": [], "files": []
    })
    folder_bucket["type1"].extend(type1)
    folder_bucket["type2"].extend(type2)
    folder_bucket["type3"].extend(type3)
    folder_bucket["files"].append(path.name)
    
    return {
        "file": str(path),
        "type1_count": len(type1),
        "type2_count": len(type2),
        "type3_count": len(type3),
        "plot": str(img_path),
        "action": "read_three_type_data"
    }


@processor(name="exit_folder_summary", priority=60, source=__file__, metadata={
    "name": "Exit Folder - Summary & Plot",
    "author": "complex_demo",
    "version": "1.0",
    "description": "On folder exit: write summary and create comprehensive plot",
})
def exit_folder_summary(path: Path, context, **cfg):
    """Write summary and comprehensive plot when exiting a folder."""
    if not path.is_dir():
        return {"skipped": True}
    
    ok, err = _ensure_packages()
    if not ok:
        return {"error": f"missing packages: {err}"}
    
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from matplotlib.figure import Figure
    from matplotlib.backends.backend_agg import FigureCanvasAgg as FigureCanvas
    
    # Get folder data
    folder_key = str(path)
    folder_data = context.get_data(["complex_demo", "folder_data", folder_key], {})
    
    type1 = folder_data.get("type1", [])
    type2 = folder_data.get("type2", [])
    type3 = folder_data.get("type3", [])
    files = folder_data.get("files", [])
    
    if not (type1 or type2 or type3):
        return {"skipped": True, "reason": "no_data"}
    
    # Get document
    doc_path = context.get_shared(["complex_demo", "doc_path"], cfg.get("doc_path", "./demo_complex_output.docx"))
    doc = Document(doc_path) if Path(doc_path).exists() else Document()
    
    # Add summary heading
    folder_name = path.name
    label = FOLDER_LABELS.get(folder_name, folder_name)
    summary_heading = doc.add_heading(f"📊 {label} - 总结", level=3)
    summary_heading.runs[0].font.color.rgb = RGBColor(0, 153, 76)
    
    # Summary text
    type1_avg = sum(type1)/len(type1) if type1 else 0
    type2_avg = sum(type2)/len(type2) if type2 else 0
    type3_avg = sum(type3)/len(type3) if type3 else 0
    
    summary_text = f"""
本目录共处理 {len(files)} 个文件。
- Type1 主要数据: {len(type1)} 个数据点
- Type2 辅助数据1: {len(type2)} 个数据点
- Type3 辅助数据2: {len(type3)} 个数据点

数据统计:
- Type1 平均值: {type1_avg:.2f}
- Type2 平均值: {type2_avg:.2f}
- Type3 平均值: {type3_avg:.2f}
"""
    doc.add_paragraph(summary_text)
    
    # Create comprehensive plot
    img_dir = Path(cfg.get("img_dir", "./demo_complex_images"))
    img_dir.mkdir(parents=True, exist_ok=True)
    img_path = img_dir / f"summary_{folder_name}.png"
    
    try:
        fig = Figure(figsize=(7, 5), dpi=120)
        ax = fig.add_subplot(111)
        
        if type1:
            ax.plot(type1, marker='o', label=f'Type1 主要数据 (n={len(type1)})', 
                   linewidth=2.5, markersize=5, alpha=0.8)
        if type2:
            ax.plot(type2, marker='s', label=f'Type2 辅助数据1 (n={len(type2)})', 
                   linewidth=2, markersize=5, alpha=0.8)
        if type3:
            ax.plot(type3, marker='^', label=f'Type3 辅助数据2 (n={len(type3)})', 
                   linewidth=2, markersize=5, alpha=0.8)
        
        ax.set_title(f"综合数据曲线 - {label}", fontsize=14, fontweight='bold')
        ax.set_xlabel("数据索引", fontsize=11)
        ax.set_ylabel("数值", fontsize=11)
        ax.legend(loc='best', fontsize=10)
        ax.grid(True, alpha=0.3, linestyle='--')
        fig.tight_layout()
        
        canvas = FigureCanvas(fig)
        canvas.draw()
        fig.savefig(img_path)
        
        # Add comprehensive plot to Word
        doc.add_paragraph("综合数据可视化：")
        doc.add_picture(str(img_path), width=Inches(6.0))
        
    except Exception as e:
        doc.add_paragraph(f"综合绘图失败: {e}")
    
    # Add separator
    doc.add_paragraph("=" * 60)
    doc.add_paragraph("")
    
    doc.save(doc_path)
    
    return {
        "folder": folder_name,
        "label": label,
        "files_processed": len(files),
        "type1_count": len(type1),
        "type2_count": len(type2),
        "type3_count": len(type3),
        "summary_plot": str(img_path),
        "action": "exit_folder_summary"
    }
