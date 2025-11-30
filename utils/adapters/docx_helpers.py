from pathlib import Path
from typing import Tuple, List, Optional, Sequence

def get_or_create_doc(doc_path: str) -> Tuple[object, Path]:
    """Return a python-docx Document and its path. Adapter to decouple docx from core utils."""
    from docx import Document
    p = Path(doc_path)
    doc = Document(p) if p.exists() else Document()
    return doc, p

def save_doc(doc, path: Path) -> None:
    doc.save(str(path))

def docx_write_text(doc, text: str = "", style: str = "Normal", align: Optional[int] = None) -> None:
    """Write a paragraph to the end of document with optional style and alignment.

    align uses python-docx alignment values: 0=LEFT, 1=CENTER, 2=RIGHT, 3=JUSTIFY.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph(text)
    try:
        p.style = style
    except Exception:
        pass
    if align is not None:
        try:
            mapping = {
                0: WD_ALIGN_PARAGRAPH.LEFT,
                1: WD_ALIGN_PARAGRAPH.CENTER,
                2: WD_ALIGN_PARAGRAPH.RIGHT,
                3: WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            p.alignment = mapping.get(align, WD_ALIGN_PARAGRAPH.LEFT)
        except Exception:
            pass

def docx_insert_table(doc, data: List[List[str]], header: Optional[List[str]] = None, caption: Optional[str] = None, style: Optional[str] = None):
    """Insert a table filled by data (2D list). Optionally add header row and caption paragraph."""
    rows = len(data) + (1 if header else 0)
    cols = len(header) if header else (len(data[0]) if data else 0)
    if rows == 0 or cols == 0:
        return None
    table = doc.add_table(rows=rows, cols=cols)
    if style:
        try:
            table.style = style
        except Exception:
            pass
    r = 0
    if header:
        hdr = table.rows[0].cells
        for j, h in enumerate(header):
            hdr[j].text = str(h)
        r = 1
    for i, row in enumerate(data):
        cells = table.rows[r + i].cells
        for j, val in enumerate(row[:cols]):
            cells[j].text = str(val)
    if caption:
        doc.add_paragraph(f"Table: {caption}")
    return table

def docx_insert_picture(doc, pic_path: Path, width_inches: Optional[float] = None, caption: Optional[str] = None):
    """Insert a picture from path with optional width (in inches) and caption."""
    from docx.shared import Inches
    if width_inches is not None:
        doc.add_picture(str(pic_path), width=Inches(width_inches))
    else:
        doc.add_picture(str(pic_path))
    if caption:
        doc.add_paragraph(f"Figure: {caption}")

def docx_insert_caption(doc, label: str = "图", text: str = "") -> None:
    """Insert a simple caption paragraph with label and optional text.

    Note: python-docx does not expose Word's dynamic caption fields directly; we emulate with text.
    """
    doc.add_paragraph(f"{label} {text}".strip())

def docx_merge_cells(table, merge_groups: Sequence[Sequence[Tuple[int, int]]]) -> None:
    """Merge specified cell groups in a table.

    merge_groups: list of groups; each group is a list of (row, col) 1-based positions.
    First cell in a group will be merged with each subsequent cell.
    """
    for group in merge_groups:
        if not group:
            continue
        r0, c0 = group[0]
        master = table.cell(r0, c0)
        for (ri, ci) in group[1:]:
            try:
                master.merge(table.cell(ri, ci))
            except Exception:
                # ignore invalid merges
                pass

def docx_table_with_caption_and_merges(
    doc,
    data: List[List[str]],
    header: Optional[List[str]] = None,
    caption: Optional[str] = None,
    style: Optional[str] = None,
    merge_groups: Optional[Sequence[Sequence[Tuple[int, int]]]] = None,
):
    """Convenience wrapper: create table, optionally merge cells, and add caption."""
    table = docx_insert_table(doc, data, header=header, caption=None, style=style)
    if table is None:
        return None
    if merge_groups:
        docx_merge_cells(table, merge_groups)
    if caption:
        docx_insert_caption(doc, label="表", text=caption)
    return table

